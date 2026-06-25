from __future__ import annotations

import argparse
import re
import textwrap
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageTemplate,
    Paragraph,
    Preformatted,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS_ROOT = ROOT / "docs"
OUTPUT_ROOT = DOCS_ROOT / "formatted_documents"
PDF_ROOT = OUTPUT_ROOT / "pdf"
DOCX_ROOT = OUTPUT_ROOT / "docx"


@dataclass
class Block:
    kind: str
    text: str
    level: int = 0


def discover_sources(docs_root: Path, output_root: Path) -> list[Path]:
    sources = [p for p in docs_root.rglob("*.md") if output_root not in p.parents]
    for extra in [ROOT / "README.md", ROOT / "RELEASE_NOTES.md"]:
        if extra.exists():
            sources.append(extra)
    return sorted({p.resolve() for p in sources})


def normalize_inline(text: str) -> str:
    text = text.replace("\ufeff", "")
    text = text.replace("\u2019", "'")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    return text


def parse_markdown(md_text: str) -> list[Block]:
    lines = md_text.splitlines()
    blocks: list[Block] = []
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                blocks.append(Block("code", "\n".join(code_lines)))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(raw.rstrip())
            i += 1
            continue

        if not stripped:
            blocks.append(Block("blank", ""))
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            blocks.append(Block("heading", normalize_inline(stripped[level:].strip()), level=level))
            i += 1
            continue

        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:-]+\|[\s|:-]*$", lines[i + 1].strip()):
            table_lines = [stripped]
            i += 2
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].rstrip())
                i += 1
            blocks.append(Block("table", "\n".join(table_lines)))
            continue

        bullet = re.match(r"^(\s*)([-*+])\s+(.*)$", raw)
        if bullet:
            indent = len(bullet.group(1)) // 2
            blocks.append(Block("bullet", normalize_inline(bullet.group(3).strip()), level=indent))
            i += 1
            continue

        numbered = re.match(r"^(\s*)(\d+)\.\s+(.*)$", raw)
        if numbered:
            indent = len(numbered.group(1)) // 2
            blocks.append(Block("number", normalize_inline(numbered.group(3).strip()), level=indent))
            i += 1
            continue

        if stripped.startswith(">"):
            blocks.append(Block("quote", normalize_inline(stripped.lstrip(">").strip())))
            i += 1
            continue

        paragraph_lines = [normalize_inline(stripped)]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith("#") or nxt.startswith("```") or nxt.startswith(">"):
                break
            if re.match(r"^(\s*)([-*+])\s+(.*)$", lines[i]) or re.match(r"^(\s*)(\d+)\.\s+(.*)$", lines[i]):
                break
            paragraph_lines.append(normalize_inline(nxt))
            i += 1
        blocks.append(Block("paragraph", " ".join(paragraph_lines)))

    if in_code and code_lines:
        blocks.append(Block("code", "\n".join(code_lines)))
    return blocks


def ensure_font_registered() -> str:
    candidates = [
        (r"C:\Windows\Fonts\Nirmala.ttc", "NirmalaUI"),
        (r"C:\Windows\Fonts\segoeui.ttf", "SegoeUI"),
        (r"C:\Windows\Fonts\arial.ttf", "Arial"),
    ]
    for path, name in candidates:
        font_path = Path(path)
        if font_path.exists():
            try:
                pdfmetrics.registerFont(TTFont(name, str(font_path)))
                return name
            except Exception:
                continue
    return "Helvetica"


def source_to_docx(source: Path, target: Path) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(11)

    for block in parse_markdown(source.read_text(encoding="utf-8-sig")):
        if block.kind == "blank":
            doc.add_paragraph("")
            continue
        if block.kind == "heading":
            level = min(block.level, 4)
            doc.add_heading(block.text, level=level)
            continue
        if block.kind == "paragraph":
            p = doc.add_paragraph()
            p.add_run(block.text)
            continue
        if block.kind == "quote":
            p = doc.add_paragraph()
            run = p.add_run(block.text)
            run.italic = True
            continue
        if block.kind in {"bullet", "number"}:
            style = "List Bullet" if block.kind == "bullet" else "List Number"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Inches(0.25 * block.level)
            p.add_run(block.text)
            continue
        if block.kind == "table":
            rows = [row.split("|") for row in block.text.splitlines()]
            cleaned_rows = [[cell.strip() for cell in row if cell.strip() != ""] for row in rows]
            if not cleaned_rows:
                continue
            cols = max(len(row) for row in cleaned_rows)
            table = doc.add_table(rows=len(cleaned_rows), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(cleaned_rows):
                for c_idx in range(cols):
                    table.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""
            continue
        if block.kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(block.text)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            continue

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)


def _set_page(doc: BaseDocTemplate) -> None:
    pass


def source_to_pdf(source: Path, target: Path, font_name: str) -> None:
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10.5,
        leading=13,
        alignment=TA_LEFT,
        spaceAfter=6,
    )
    body_bold = ParagraphStyle(
        "BodyBold",
        parent=body,
        fontName=font_name,
        leading=13,
    )
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName=font_name, fontSize=18, leading=22, spaceAfter=8)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName=font_name, fontSize=15, leading=18, spaceAfter=7)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontName=font_name, fontSize=13, leading=16, spaceAfter=6)
    code_style = ParagraphStyle(
        "Code",
        parent=styles["Code"],
        fontName=font_name,
        fontSize=9,
        leading=11,
        backColor=colors.whitesmoke,
        borderPadding=6,
        leftIndent=6,
        rightIndent=6,
        spaceAfter=6,
    )

    frame = Frame(inch * 0.7, inch * 0.65, A4[0] - inch * 1.4, A4[1] - inch * 1.3, id="body")
    doc = BaseDocTemplate(str(target), pagesize=A4, leftMargin=0.7 * inch, rightMargin=0.7 * inch, topMargin=0.65 * inch, bottomMargin=0.65 * inch)
    doc.addPageTemplates([PageTemplate(id="page", frames=[frame])])

    story: list[object] = []
    for block in parse_markdown(source.read_text(encoding="utf-8-sig")):
        if block.kind == "blank":
            story.append(Spacer(1, 0.08 * inch))
            continue
        if block.kind == "heading":
            if block.level <= 1:
                story.append(Paragraph(block.text, h1))
            elif block.level == 2:
                story.append(Paragraph(block.text, h2))
            else:
                story.append(Paragraph(block.text, h3))
            continue
        if block.kind == "paragraph":
            story.append(Paragraph(block.text, body))
            continue
        if block.kind == "quote":
            story.append(Paragraph(f"&gt; {block.text}", body))
            continue
        if block.kind == "bullet":
            indent = "&nbsp;" * (block.level * 4)
            story.append(Paragraph(f"{indent}• {block.text}", body))
            continue
        if block.kind == "number":
            indent = "&nbsp;" * (block.level * 4)
            story.append(Paragraph(f"{indent}1. {block.text}", body))
            continue
        if block.kind == "table":
            rows = [row.split("|") for row in block.text.splitlines()]
            cleaned_rows = [[cell.strip() for cell in row if cell.strip() != ""] for row in rows]
            if not cleaned_rows:
                continue
            cols = max(len(row) for row in cleaned_rows)
            tbl = Table(
                [[cell for cell in row] + [""] * (cols - len(row)) for row in cleaned_rows],
                hAlign="LEFT",
                repeatRows=1,
            )
            tbl.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("FONTNAME", (0, 0), (-1, -1), font_name),
                        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                        ("LEADING", (0, 0), (-1, -1), 10),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(tbl)
            continue
        if block.kind == "code":
            story.append(Preformatted(block.text, code_style))
            continue

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)


def export_all(docs_root: Path, output_root: Path) -> None:
    pdf_root = output_root / "pdf"
    docx_root = output_root / "docx"
    font_name = ensure_font_registered()
    sources = discover_sources(docs_root, output_root)
    for source in sources:
        relative = source.relative_to(docs_root) if docs_root in source.parents else Path(source.name)
        pdf_target = (pdf_root / relative).with_suffix(".pdf")
        docx_target = (docx_root / relative).with_suffix(".docx")
        source_to_pdf(source, pdf_target, font_name)
        source_to_docx(source, docx_target)
        print(f"Exported {source} -> {pdf_target} and {docx_target}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Export markdown docs to PDF and DOCX using standard libraries.")
    parser.add_argument("--source-root", type=Path, default=DOCS_ROOT)
    parser.add_argument("--output-root", type=Path, default=OUTPUT_ROOT)
    args = parser.parse_args()

    export_all(args.source_root.resolve(), args.output_root.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
